"""
Модуль для парсинга word файл.

Код модуля обрабатывает word файл и получает из него нужные данные,
подлежащие передаче для дальнейшей работе по преобразованию и сравнению.
Необходимые данные становятся свойствами объекта класса MainDataGetter.
Содержит классы и методы, направленные на извлечение данных.

Класс:
    - MainCollector:
        Верхне-уровневый класс для сбора данных.

    - CollectorFromTable:
        Класс для получения информации из таблиц word файла.


Функции:
    - prepare_for_work_with_tables:
        Подготовить word файл для работы с таблицами - преобразовать в нужный формат.

    - find_out_type_of_table:
        Определить тип таблицы по содержанию первой строки, сопоставив с паттернами.

    - get_one_string_from_file:
        Прочитать word файл, вернуть одну большую строку.

    - get_main_information:
        Найти и вернуть номер и дату протокола.

    - collect_prod_control_data:
        Найти и вернуть данные, относящиеся к измерениям производственного контроля

Пример использования:
    data_getter = MainCollector(FILE)
    data_getter.collect_all_data() - собрать все данные
    data_getter.main_number - атрибут
    data_getter.main_date - атрибут
    data_getter.data_from_tables - атрибут
    data_getter.prod_control_data - атрибут

"""
import itertools
import re

from docx import Document
from docx.table import Table
from docx2txt import docx2txt

from league_sert.constants import (PattNumbDate, ProdControlPatt, TypesOfTable,
                                   WRONG_PARTS_IN_ROW, ConvertValueTypes, ComparTypes, WordsPatterns, TableWords,
                                   TABLE_WORDS_PATTERN)
from league_sert.data_preparation.exceptions import TypeOfTableError
from league_sert.data_preparation.extract_numb_and_date import get_main_numb_and_date


def check_wrong_parts_in_row(one_row: str, wrong_items) -> bool:
    """ Проверить строку из таблицы на валидность значений.
     Получает одну строку, в которой сохранены все строки из ячеек строки
     таблицы. Смотрит есть ли в ней слова, которые указывают на то,
     что эта строка не нужна. """
    for pat in wrong_items:
        if re.search(pat, one_row):
            return False
    return True


def get_text_with_superscripts(cell):
    """ Получить текст ячейки, сохраняя при этом числа со степенью. """
    text = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.superscript:
                if run.text == 'б':
                    text += f"{'⁰¹²³⁴⁵⁶⁷⁸⁹'[int(6)]}"
                else:
                    for i in run.text:
                        if i.isdigit():
                            text += f"{'⁰¹²³⁴⁵⁶⁷⁸⁹'[int(i)]}"
                        else:
                            text += i
            else:
                text += run.text
    return text


def check_table_only_symbols(table_: Table):
    """ Проверить, что таблица валидная, а не состоит только из символов.
    Весь текст из ячеек собираем в один сплошной фрагмент.
    Затем ищем в таблице определенные слова, свидетельствующие о том, что
    это таблица, а не набор символов. """

    row_count = 0
    solid_text = ''
    for row in table_.rows:
        row_count += 1
        for cell_ in row.cells:
            solid_text += cell_.text

    matches = re.search(TABLE_WORDS_PATTERN, solid_text, re.IGNORECASE)
    return True if matches else False


def find_out_table_type(first_row_cells: list, table_: Table) -> str | None:
    """ Определить тип таблицы по содержанию первой строки в таблице.
    Возможные типы - это основная, данные о пробе, результаты исследования, данные
    о средствах измерения производственного контроля, сам производственный контроль.

    :param first_row_cells: Список всех ячеек первой строки таблицы.
    :return: Либо строкой тип таблицы, либо None. """

    # Текст из ячеек первой строки таблицы.
    cells_text = ''
    for cell_ in first_row_cells:
        cells_text += (cell_.text + ' ')

    # Перебираем паттерны таблиц и сравниваем с содержимым первой строки
    for pattern in TypesOfTable:
        if re.search(pattern.value, cells_text, re.IGNORECASE):
            return pattern.name  # Тип таблицы
        if re.search(r'Данная проба по исследованным ', cells_text, re.IGNORECASE):
            return None

    # Если не определен валидный тип таблицы.
    # Проверяем полноценная ли и самодостаточная это таблица.
    if len(first_row_cells) <= 2:  # Смотрим сколько ячеек в строке.
        return None

    # Здесь вызов метода на проверку, что это не таблица с символами.
    if not check_table_only_symbols(table_):
        return None

    # Проверяем не является ли это продолжением таблицы с результатами.
    # Если является, то возвращаем тип таблицы результаты RESULTS.
    for pattern in ConvertValueTypes:
        if pattern == ConvertValueTypes.NONE:
            continue
        match_ = re.search(pattern.value, first_row_cells[1].text)
        match_2 = re.search(pattern.value, first_row_cells[2].text)
        if match_ or match_2:
            return TypesOfTable.RESULTS.name

    # Для отладки, если тип таблицы не определен, то поднять исключение.
    raise TypeOfTableError(first_row_cells)


def delete_row(table_: Table, row_numb: int) -> Table:
    """
    Удаляет первую строку из переданной таблицы и возвращает изменённую таблицу.

    :param table: Таблица, из которой нужно удалить первую строку.
    :return: Изменённая таблица без первой строки.
    """
    # Проверяем, что в таблице больше одной строки
    if len(table_.rows) <= 1:
        raise ValueError("Таблица должна содержать более одной строки, "
                         "чтобы удалять первую строку.")
    # Извлекаем индекс первой строки и удаляем её
    table_._tbl.remove(table_.rows[row_numb]._tr)
    return table_


def remove_wrong_rows_in_table(tables, patterns=WRONG_PARTS_IN_ROW):
    """ Удалить из таблицы строки, которые содержат определенные слова. """

    # Компиляция регулярных выражений
    compiled_patterns = [re.compile(pattern) for pattern in patterns]

    # Пройтись по всем таблицам в документе
    for table in tables:
        rows_to_delete = []

        # Пройтись по всем строкам в таблице
        for row_idx, row in enumerate(table.rows):
            row_text = ' '.join(cell.text for cell in row.cells)

            # Проверка на совпадение строки с любым шаблоном
            if any(pattern.search(row_text) for pattern in compiled_patterns):
                rows_to_delete.append(row_idx)

        # Удалить отмеченные строки (в обратном порядке чтобы индексы не смещались)
        for row_idx in reversed(rows_to_delete):
            table._element.remove(table.rows[row_idx]._element)

    # Возвращаем модифицированный документ
    return tables


def _replace_to_6_in_text(match):
    """ Функция для замены <б> на <6> в re.sub """
    digit = match.group(1)  # Получаем цифру, если она есть
    if digit:  # Если цифра найдена
        return f"«{digit}6»"
    else:  # Если цифра не найдена
        return "«6»"  # Просто возвращаем 6


def collect_prod_control_data(text) -> dict:
    """ Найти и вернуть данные, относящиеся к измерениям
    производственного контроля. """

    try:
        # Ищем начало данных о производственном контроле
        substr_start = re.search(ProdControlPatt.START_SUBSTR.value, text).start()

        # Берем запас на случай улетевшей ранее даты протокола.
        text = text[substr_start - 100:]

        # Заменяем возможные ошибки при распознавании 'б' на '6'.
        text = re.sub(r'«(\d)?\s*б»', _replace_to_6_in_text, text)
        text = re.sub(r'«II»', "«11»", text)
        text = re.sub(r'«И»', "«11»", text)

        # Ищем номер протокола.
        match_number_of_protocol = re.search(PattNumbDate.NUMBER.value, text)
        number_of_protocol = match_number_of_protocol.group(1)

        # Ищем дату протокола
        match_date = re.search(ProdControlPatt.DATE.value, text)
        date_of_protocol = (match_date.group(1) + ' ' + match_date.group(2) + ' ' + match_date.group(3))

        # Ищем дату измерений.
        match = re.search(ProdControlPatt.NUMB_AND_DATE_OF_MEASUR.value, text)
        date_of_measurement = match.group(1)

        # Ищем место измерений.
        match = re.search(ProdControlPatt.PLACE_OF_MEASUR.value, text)
        place_of_measurement = match.group(1).strip()

        # Ищем содержание заключения.
        inner_conclusion = re.search(ProdControlPatt.INNER_PART_OF_CONCLUSION.value, text)
        if inner_conclusion:
            inner_conclusion = inner_conclusion.group().strip()
        else:
            inner_conclusion = None

        return {'number_of_protocol': number_of_protocol,
                'date_of_protocol': date_of_protocol,
                'act': date_of_measurement,
                'place_of_measurement': place_of_measurement,
                'inner_conclusion': inner_conclusion}

    except AttributeError:
        print('Таблицы производственного контроля не найдено')


class CollectorFromTable:
    """ Класс для получения информации из таблиц word файла. """

    def __init__(self, document_with_tables):
        self.all_tables_in_file: list = document_with_tables.tables
        self.converted_data_from_tables = {}
        self.current_table: Table
        self.key_of_current_table: str
        self.value_of_current_table: list | dict

    def process_two_columns_table(self):
        """ Преобразовать таблицу, состоящую из двух колонок в словарь,
         где ключ значение из первой колонки, а значение из второй колонки."""

        self.value_of_current_table = {}
        for _, row in enumerate(self.current_table.rows):

            cells = row.cells
            for _ in cells:
                key_of_row = cells[0].text.strip()
                key_of_row = re.sub(r'11(\w+)', r'Н\1', key_of_row)
                key_of_row = key_of_row.strip('.').strip()
                value_of_row = cells[1].text.strip()
                self.value_of_current_table[key_of_row] = value_of_row
        self.converted_data_from_tables[self.key_of_current_table] = self.value_of_current_table

    def process_more_then_two_cols(self):
        """ Преобразовать таблицу, состоящую из более чем двух колонок
        в список со вложенными списками. Каждый вложенный список это значения
        из одной строки таблицы. """

        # Валидные значения из строки таблицы.
        self.value_of_current_table = []
        for _, row in enumerate(self.current_table.rows):
            one_row = []
            cells = row.cells

            # Формируем список из элементов одной строки в таблице.
            for cell in cells:
                cell_text = get_text_with_superscripts(cell)
                one_row.append(cell_text)

            # Проверяем есть ли в строке недопустимые значения.
            if check_wrong_parts_in_row(" ".join(one_row), WRONG_PARTS_IN_ROW):
                self.value_of_current_table.append(one_row)

        self.converted_data_from_tables[self.key_of_current_table] = self.value_of_current_table

    def extract_data_from_all_tables(self):
        """ Собрать в словарь, данные из всех таблиц word файла,
         ключом будет tuple, где первое значение порядковый номер таблицы в документе
         второе тип таблицы. Значением в итоговом словаре для таблицы
         будут данные либо в виде списка со вложенными списками,
         либо в виде словаря. """

        # Убрать строки, которые не нужны для дальнейшей работы
        self.all_tables_in_file = remove_wrong_rows_in_table(self.all_tables_in_file)

        # Перебираем таблицы.
        for table_number, table_ in enumerate(self.all_tables_in_file):

            self.current_table = table_  # Текущая таблица

            # Проверка, что есть более одной 1 строки и более 1 колонки.
            if len(table_.columns) <= 1:
                continue

            # Ячейки из первой строки таблицы, по ним поймем, что за таблица.
            first_row_cells = self.current_table.row_cells(0)

            if not first_row_cells:
                continue

            cells_text = ''
            for cell_ in first_row_cells:
                cells_text += (cell_.text + ' ')
            if set(cells_text) == {' '}:
                first_row_cells = self.current_table.row_cells(1)

            # Определить тип таблицы.
            type_of_table = find_out_table_type(first_row_cells, self.current_table)
            # Строку выше засунуть в блок, если вылезла ошибка, то

            if type_of_table is None:
                continue

            # Сформировать ключ для таблицы в общих данных.
            self.key_of_current_table = (table_number, type_of_table)

            # Выбрать метод обработки данных в таблице в зависимости
            # от типа таблицы количества колонок.
            if type_of_table in {TypesOfTable.MAIN.name, TypesOfTable.SAMPLE.name}:
                self.process_two_columns_table()
            if type_of_table in {TypesOfTable.RESULTS.name, TypesOfTable.PROD_CONTROL.name}:
                self.process_more_then_two_cols()


class MainCollector:
    """ Верхне-уровневый класс для получения данных из файла. """

    def __init__(self, file_path):
        self.file = file_path  # Исходный файл.
        self.tables_from_file = None  # Таблицы из файла для обработки.
        self.text = None  # Сплошной текст из файла для обработки.
        self.data_from_tables = None  # Результат - данные по таблицам.
        self.main_number = None  # Основной номер.
        self.main_date = None  # Основная дата.
        self.prod_control_data = None  # Данные по производственному контролю.

    def get_data_from_tables(self):
        """ Собрать все данные из таблиц. """
        table_collector = CollectorFromTable(self.prepare_for_work_with_tables())
        table_collector.extract_data_from_all_tables()
        self.data_from_tables = table_collector.converted_data_from_tables

    def get_string(self):
        """ Получить строку для поиска строк, вне таблиц"""
        if not self.text:
            self.text = docx2txt.process(self.file)
        return self.text

    def prepare_for_work_with_tables(self):
        """ Подготовить файл для извлечения данных из таблиц. """
        if not self.tables_from_file:
            self.tables_from_file = Document(self.file)
        return self.tables_from_file

    def get_main_information(self):
        """ Получить основные данные из файла. """
        if not self.main_number or not self.main_date:
            self.main_number, self.main_date = get_main_numb_and_date(self.get_string(), self.file)

    def get_prod_control_data(self):
        """ Получить из файла данные о производственном контроле. """
        if not self.prod_control_data:
            self.prod_control_data = collect_prod_control_data(self.get_string())

    def collect_all_data(self):
        """ Собрать все данные из word файла. """
        self.get_main_information()
        self.get_data_from_tables()
        self.get_prod_control_data()

    def merge_prod_control(self):
        """ Объединить воедино данные по производственному контролю. """
        for key, value in self.data_from_tables.items():
            if key[1] == 'PROD_CONTROL':
                updated = {**self.prod_control_data, 'indicators': [value]}
                self.data_from_tables[key] = updated
                continue
