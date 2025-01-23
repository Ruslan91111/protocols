"""
Модуль для конвертации PDF файлов в Word и проверки наличия
номеров и дат протоколов в конвертированных файлах.

Этот модуль предоставляет функциональность для автоматического
конвертирования всех PDF файлов в указанной директории в формат Word
с помощью "кликера" и программы FineReader и последующей проверки данных в каждом
конвертированном документе. Модуль включает в себя два основных
класса: `FileConverter` для выполнения самой конвертации и
`WordChecker` для проверки содержимого файлов на наличие специфических
данных, таких как номера и даты протоколов.

Классы:

    - `FileConverter`: Класс, реализующий логику конвертации одного
      документа из PDF в Word. Он позволяет выбрать исходный PDF файл
      и задать целевую директорию для сохранения результата в формате Word.

    - `WordChecker`: Класс, предоставляющий методы для проверки
      конвертированного файла в формате Word. Он ищет в документе
      номера основного протокола и даты протоколов производственного контроля.

Функции:

    - `check_word_file(file_path: str) -> tuple | False`:
      Проверяет указанный Word файл на наличие номеров и дат протоколов.
      Возвращает кортеж с найденными значениями, если проверки успешны,
      либо False, если необходимые данные не найдены.

    - `convert_a_file(file_path: str, dir_with_pdf_files: str, dir_for_word_files: str) -> None`:
      Конвертирует один указанный PDF файл в формат Word и сохраняет
      его в заданную директорию. Осуществляет взаимодействие с
      внешним "кликером" для выполнения конвертации.

    - `convert_all_pdf_in_dir_to_word(dir_with_pdf_files: str) -> None`:
      Запрашивает у пользователя путь к директории с PDF файлами. Создаёт
      директорию для сохранения Word файлов и запускает процесс
      конвертации всех PDF файлов в указанной директории.
      После конвертации, выполняет проверку каждого Word файла
      на наличие номеров и дат протоколов, используя `WordChecker`.

Пример использования:
    convert_all_pdf_in_dir_to_word(dir_with_pdf_files)

"""
import os
import re
import shutil
import time
from pathlib import Path

from docx import Document
from docx2txt import docx2txt

from league_sert.constants import (FRScreens,
                                   FINE_READER_PROCESS,
                                   PattNumbDate,
                                   ProdControlPatt)

from league_sert.data_preparation.extract_numb_and_date import get_main_numb_and_date

from conversion.exceptions import (ScreenshotNotFoundError,
                                   NumbDateNotFoundError,
                                   PageNeedToRMError)

from conversion.remove_back_elements import rm_stamps_from_pdf

from conversion.files_and_proc_works import (return_or_create_dir,
                                             fetch_files_for_conversion,
                                             is_file_in_dir)

from conversion.screen_work import (click_scr,
                                    launch_desktop_app,
                                    wait_fr_app_loading,
                                    click_convert_main_menu,
                                    input_filename,
                                    is_in_convert_to_word_section,
                                    uncheck_save_image,
                                    uncheck_open_doc,
                                    click_convert_blue_button,
                                    handle_fr_warning,
                                    check_save_image, wait_scr)


def input_filename_for_conversion(file, screenshot):
    """ Ввести в поле путь к файлу, нажать enter и дождаться появления
     переданного скриншота, чтобы понять, что ввод файла прошел успешно."""
    time.sleep(0.5)
    # Вводим путь к файлу.
    input_filename(file)
    try:
        # Проверяем появился ли скриншот, который свидетельствует об успешном вводе пути.
        wait_scr(screenshot)

    # В случае ошибки повторяем ввод пути к файлу.
    except ScreenshotNotFoundError:
        input_filename(file)


class FileConverter:
    """ Класс конвертер одного документа."""

    def __init__(self, file: str, dir_pdf_files: str, dir_word_files: str):
        self.file = file  # Файл для конвертации.
        self.dir_pdf_files = dir_pdf_files  # Директория с не конвертированными ПДФ файлами.
        self.dir_word_files = dir_word_files  # Директория для конвертированных word файлов.
        self.file_pdf_original = os.path.join(dir_pdf_files, file) + '.pdf'  # Путь к ПДФ файлу.
        self.file_word_original = os.path.join(dir_word_files, file) + '.docx'  # Путь к word файлу.
        self.temp_pdf_file = None  # Файл, без печатей и подписей синего цвета.
        self.temp_word_file = None  # Word для конвертации PDF, без печатей.
        self.numbs_and_dates = None  # Номер и дата протокола

    def input_pdf_origin(self):
        """ Ввести путь к оригинальному ПДФ файлу. """
        input_filename_for_conversion(
            self.file_pdf_original, FRScreens.CONVERT_TO_WORD_INNER_BUTTON.value)

    def input_path_to_word(self):
        """ Ввести путь к word файлу, в который будет конвертирован ПДФ. """
        input_filename_for_conversion(
            self.file_word_original, FRScreens.PROCESS_OF_CONVERSION.value)

    def convert_file(self):
        """ Выполнить действия по конвертированию одного файла. """
        click_convert_main_menu()
        self.input_pdf_origin()
        check_save_image()
        click_convert_blue_button()
        uncheck_open_doc()
        self.input_path_to_word()
        is_file_in_dir(self.dir_word_files, self.file + '.docx')
        handle_fr_warning()
        click_scr(FRScreens.BUTTON_CANCEL.value)

    def check_numbs_and_date(self):
        """ Проверить есть ли в протоколе номера и даты. """
        self.numbs_and_dates = check_word_file(self.file_word_original)
        if self.numbs_and_dates:
            return True

    def convert_file_without_stamps(self):
        """ Набор действий по повторной конвертации
        pdf файла без печатей. """

        # Наименования для временных word и pdf файлов.
        self.temp_pdf_file = self.file_pdf_original.replace('.pdf', '_no_stamps.pdf')
        self.temp_word_file = self.file_word_original.replace('.docx', '_temp.docx')

        # Удалить печати из PDF файла.
        rm_stamps_from_pdf(self.file_pdf_original, self.temp_pdf_file)

        # Навигация по меню FineReader.
        click_convert_main_menu()
        input_filename(self.temp_pdf_file)
        uncheck_save_image()  # Убрать галку с "сохранить изображения".
        click_convert_blue_button()
        uncheck_open_doc()  # Убрать галку с "открыть файл".
        input_filename(self.temp_word_file)
        is_file_in_dir(self.dir_word_files,
                       self.temp_word_file[self.temp_word_file.rfind('\\') + 1:])

        # Проверить наличие предупреждения, возникающего при конвертации.
        handle_fr_warning()
        click_scr(FRScreens.BUTTON_CANCEL.value)

    def add_numbs_and_dates_to_file(self):
        """ Получить из word файла без печатей номера и даты протоколов
        и дописать их в оригинальный word файл."""

        # Пытаемся получить из файла без печатей номера и даты протоколов.
        try:
            (main_date,
             main_numb,
             prod_contr_numb,
             prod_contr_date) = check_word_file(self.temp_word_file)

            #  Действия по дописыванию номеров и дат в файл Word протокола.
            doc = Document(self.file_word_original)
            doc.add_paragraph(f'Протокол испытаний № {main_numb} от {main_date}')
            doc.add_paragraph(f'ПРОТОКОЛ № {prod_contr_numb} измерений '
                              f'по производственному контролю от {prod_contr_date}')
            doc.save(self.file_word_original)
            print(f'В файл {self.file_word_original} дописаны номера и даты протоколов.')

        except (AttributeError, NumbDateNotFoundError):
            print('Нет номеров и дат протоколов.')

        except Exception as e:
            print(e)

        # Удаление временных файлов.
        os.remove(self.temp_pdf_file)
        os.remove(self.temp_word_file)

    def convert_and_check_numbs_and_dates(self):
        """ Весь процесс конвертации файла, проверки на номера. """
        # Конвертировать файл.
        self.convert_file()
        # Проверить файл на наличие номеров и дат.
        numbs_and_dates = self.check_numbs_and_date()
        if numbs_and_dates:
            return

        # Действия по удалению печатей из файла, конвертации и добавлению в файл.
        self.convert_file_without_stamps()
        self.add_numbs_and_dates_to_file()


def convert_a_file(file_path: str, dir_with_pdf_files: str, dir_for_word_files: str) -> None:
    """ Конвертировать один файл. Функция для использования
    класса FileConverter."""
    converter = FileConverter(file_path, dir_with_pdf_files, dir_for_word_files)
    converter.convert_and_check_numbs_and_dates()


class WordChecker:
    """ Класс для проверки конвертированного файла формата word
    на наличие номера и даты основного протокола и протокола
    производственного контроля. """

    def __init__(self, file_path):
        self.file_path = file_path
        self.text = docx2txt.process(self.file_path)
        self.main_numb = None
        self.main_date = None
        self.prod_contr_numb = None
        self.prod_contr_date = None

    def get_main_numb_and_date(self):
        """ Проверить наличие номера и даты основного протокола. """
        self.main_numb, self.main_date = get_main_numb_and_date(self.text, self.file_path)

    def get_prod_contr_numb_and_date(self):
        """ Проверить наличие номера и даты протокола производственного контроля. """

        try:
            substr_start = re.search(ProdControlPatt.START_SUBSTR.value, self.text).start()
            prod_control_text = self.text[substr_start:]
            self.prod_contr_numb = re.search(PattNumbDate.NUMBER.value, prod_control_text).group(1)
            match = re.search(ProdControlPatt.DATE.value, prod_control_text)
            self.prod_contr_date = match.group(1) + ' ' + match.group(2) + ' ' + match.group(3)

        except Exception as e:
            print(e)


def check_word_file(file_path: str) -> tuple | bool:
    """ Проверить Word файл на наличие номеров и дат протоколов.
    Функция для использования класса WordChecker. """

    word_checker = WordChecker(file_path)
    # Получаем основные номер и дату протокола.
    word_checker.get_main_numb_and_date()
    # Получаем номер и дату протокола производственного контроля.
    word_checker.get_prod_contr_numb_and_date()

    # Если в протоколе найдены и даты и номера, то вернуть их кортежем.
    if (word_checker.main_date and word_checker.main_numb
            and word_checker.prod_contr_numb and word_checker.prod_contr_date):
        return (word_checker.main_date, word_checker.main_numb,
                word_checker.prod_contr_numb, word_checker.prod_contr_date)

    return False


def move_unconverted_file(file, dir_with_pdf_files):
    """ Если при конвертации файла произошла ошибка, то файл,
    на котором произошла ошибка переместить в отдельную директорию. """
    # Если в название файла нет расширения, то добавить.
    file = file + '.pdf' if file[-4:] != 'pdf' else file
    # Директория для неконвертируемых файлов.
    unconverted_dir = Path(dir_with_pdf_files) / 'unconverted_files'
    unconverted_dir.mkdir(parents=True, exist_ok=True)
    # Перемещаем файл.
    original_path = Path(dir_with_pdf_files) / file
    destination_path = Path(unconverted_dir) / file
    shutil.move(original_path, destination_path)


def move_not_renamed_file(file, dir_with_word_files, dir_with_pdf_files):
    """ Если при конвертации файла произошла ошибка, то файл,
    на котором произошла ошибка переместить в отдельную директорию. """

    # Если в название файла нет расширения, то добавить.
    if file.endswith('docx'):
        pdf_file = file.replace('.docx', '.pdf')
        word_file = file
    elif file.endswith('.pdf'):
        word_file = file.replace('.pdf', '.docx')
        pdf_file = file
    else:
        pdf_file = file + '.pdf'
        word_file = file + '.docx'

    # Перемещаем word файл.
    # Директория для не переименованных файлов.
    not_renamed_word_dir = Path(dir_with_word_files) / 'not_renamed_files'
    not_renamed_word_dir.mkdir(parents=True, exist_ok=True)
    original_path = Path(dir_with_word_files) / word_file
    destination_path = Path(not_renamed_word_dir) / word_file
    shutil.move(original_path, destination_path)  # Перемещаем файл.

    # Перемещаем pdf файл.
    not_renamed_pdf_dir = Path(dir_with_pdf_files) / 'not_renamed_files'
    not_renamed_pdf_dir.mkdir(parents=True, exist_ok=True)
    original_path = Path(dir_with_pdf_files) / pdf_file
    destination_path = Path(not_renamed_pdf_dir) / pdf_file
    shutil.move(original_path, destination_path)


def convert_all_pdf_in_dir_to_word(dir_with_pdf_files: str) -> None:
    """ Получить от пользователя путь к директории с PDF файлами создать директорию,
    в которую будут сохранены файлы в формате word после конвертации. """

    # Директория для Word файлов.
    dir_for_word_files = return_or_create_dir(dir_with_pdf_files + '\\word_files\\')
    # Переменная для файлов, нуждающихся в конвертации.
    files_for_convert = True

    # Пока есть файлы для конвертации.
    while files_for_convert:

        # Перечень файлов, которые еще не конвертированы.
        files_for_convert = fetch_files_for_conversion(
            dir_with_pdf_files, dir_for_word_files)

        if not files_for_convert:
            print("Отсутствуют файлы для конвертации.")
            break

        # Запустить FineReader.
        launch_desktop_app(
            FINE_READER_PROCESS, FRScreens.PANEL_ICON.value,
            FRScreens.DESKTOP_ICON.value, 0.8)
        # Ждем загрузки приложения FineReader.
        wait_fr_app_loading()
        # Проверяем, что находимся в нужном разделе приложения.
        is_in_convert_to_word_section()

        # Перебираем файлы в перечне файлов, нуждающихся в конвертации.
        for file in files_for_convert:

            try:
                # Конвертируем файл.
                convert_a_file(file, dir_with_pdf_files, dir_for_word_files)

            # Если не найдены дата или номер протокола, то просто продолжить работу,
            # так как файл сконвертирован, и word файл есть в директории.
            except NumbDateNotFoundError:
                continue

            # Если не найден, какой-либо из скриншотов, то выйти из внутреннего цикла
            # и запустить новую итерацию попытки, начиная с запуска FineReader.
            except ScreenshotNotFoundError:
                break

            # Если есть страница, которую нужно удалить, то переместить файл
            # в отдельную директорию.
            except PageNeedToRMError:
                move_unconverted_file(file, dir_with_pdf_files)
                # Закрыть предупреждение.
                try:
                    click_scr(FRScreens.SHUT_WARNING_BLUE_FRAME.value)
                except ScreenshotNotFoundError:
                    click_scr(FRScreens.SHUT_WARNING.value)


if __name__ == '__main__':
    dir_path = r'C:\Users\RIMinullin\Desktop\2024'
    convert_all_pdf_in_dir_to_word(dir_path)
