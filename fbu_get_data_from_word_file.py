"""
"""
import json
import re
from enum import Enum

from docx import Document
import docx2txt

from fbu_comparison_results_and_norms import compare_result_and_norms


# Множество для хранения заголовков таблицы, для пропуска при высчитывании соответствия нормам.
KEYS_IN_INDICATORS_FOR_PASS = {'Токсичные элементы, мг/кг:',
                               'Пестициды, мг/кг:', }


# Ключи для поиска на первой странице протокола.
REQUIRED_KEYS_FOR_PARSING_FIRST_PAGE = (
    'Место отбора проб',
    'Дата и время отбора проб',
    'Сопроводительные документы',
    'Группа продукции',
    'Наименование продукции',
    'Дата производства продукции',
    'Производитель (фирма, предприятие, организация)',
    'Дата проведения исследований'
)


class FbuPatternsForKeys(Enum):
    """ Паттерны для изъятия данных с протоколов ФБУ."""
    number_and_date: str = r'(\bпротокол\b[\s\w]+№)([\s\d-]*)от([\w\.\s]*\d{4})'
    accompanying_documents: str = r'(\bсопроводительный\b[\w\s]+:\s*)([\w\.\s№-]+.\d{4})'
    name_of_test_sample: str = r'(Наименование образца испытаний[\*\:]*)([\w%\s\.]+)(Изготовитель)'
    manufacturer: str = r'(Изготовитель[\*\:]*)([\w%\s"\.]+)(Место)'
    date_of_manufacture: str = r'(Дата изготовления[\*\:]*)(\d{2}.\d{2}.\d{4})'
    date_of_testing: str = r'(Дата проведения испытаний[:\s]*)(с \d{2}.\d{2}.\d{4} по \d{2}.\d{2}.\d{4})'


class WordFileParser:
    """
    Класс для изъятия данных из Word файла.
    """

    def __init__(self, input_word_file: str):
        self.word_file = input_word_file
        self.all_text_from_file = self.get_one_string_from_word_file()
        self.document_with_tables = self.convert_word_file_to_docx_document()
        self.data = {}

    def get_one_string_from_word_file(self) -> str:
        """
        Прочитать word файл, вернуть одну большую строку.
        """
        text = docx2txt.process(self.word_file)
        return text

    def convert_word_file_to_docx_document(self) -> Document:
        """
        Прочитать word файл, преобразовать для работы с таблицами в
        объект класса Document.
         """
        document = Document(self.word_file)
        return document

    def get_data_by_keys(self) -> None:
        """
        Собрать данные из word файла и сохранить ко ключам в self.data.
        """
        for key in FbuPatternsForKeys:
            match = re.search(key.value, self.all_text_from_file, flags=re.IGNORECASE)
            self.data[key.name] = match.group(2).strip() if match else ''

    def get_text_with_superscripts(self, cell):
        """
        Получить текст ячейки, обрабатывая при этом числа со степенью.
        """
        text = ""

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.superscript:
                    text += f"^{run.text}"
                else:
                    text += run.text

        # Замена формата "^<digit>" на соответствующие символы степени.
        text = re.sub(r'(\d+)\^(\d+)',
                      lambda x: x.group(1) + '⁰¹²³⁴⁵⁶⁷⁸⁹'[int(x.group(2))],
                      text)
        return text

    def get_indicators_from_word_file(self) -> None:
        """
        Пройтись по таблицам с результатами исследований и собрать
        данные по ним и сохранить в поле 'indicators'.
        Результат будет сохранен в виде списка словарей, где ключ - наименование
        показателя, значение - tuple из нормы и результата исследования.
        """

        all_tables_in_file: list = self.document_with_tables.tables
        result_list = []

        for table in all_tables_in_file:  # Цикл по таблицам.
            if len(table.rows) > 1:  # Проверка на наличие в таблице более двух строк.
                indicators = {}

                # Цикл по строкам в таблице, начиная со второй.
                for _, row in enumerate(table.rows[1:], start=1):
                    cells = row.cells
                    indicator_name = self.get_text_with_superscripts(cells[0]).strip()
                    result = self.get_text_with_superscripts(cells[3]).strip()
                    standards = self.get_text_with_superscripts(cells[2]).strip()
                    indicators[indicator_name] = (result, standards)

                result_list.append(indicators)

        self.data['indicators']: list[dict] = result_list

    def add_a_conclusion_about_compliance_of_norms(self):
        """
        Добавить в словарь выводы о соответствие каждого результатов исследований нормам,
        а также сделать общий вывод для протокола имеются ли в нем несоответствия нормам.
        """

        violations = False   # Переменная есть ли нарушения в показателях.

        # Перебираем словари в 'показателях'.
        for number, indicator_dict in enumerate(self.data['indicators']):
            # key - наименование показателя, value - tuple: (результат, норма для показателя)
            for key, value in indicator_dict.items():

                # Пропуск ключей, по которым не нужно проводить сравнение.
                if key in KEYS_IN_INDICATORS_FOR_PASS:
                    new_value = value + ('-',)

                else:
                    # Получаем bool соответствует ли результат норме.
                    complies_with_standards = compare_result_and_norms(value[0], value[1])
                    if not complies_with_standards:
                        violations = True
                    new_value = value + (complies_with_standards,)

                # Добавляем вывод о соответствии в словарь
                self.data['indicators'][number][key] = new_value

        # Добавляем в словарь с данными по всему протоколу ключ
        # с bool есть ли в протоколе нарушение норм.
        self.data['violation_of_norms'] = 'Не соответствуют' if violations else 'Соответствуют'

    def get_all_required_data_from_word_file(self):
        """ Основной метод - направленный на изъятие данных из Word файла."""
        self.get_data_by_keys()  # Получить основные данные по протоколу
        self.get_indicators_from_word_file()  # Добавить в словарь показатели.
        self.add_a_conclusion_about_compliance_of_norms()  # Вывод о соответствии нормам
        # self.data['indicators'] = json.dumps(self.data['indicators'])  # преобразовать в json


if __name__ == '__main__':

    wp=WordFileParser(r'C:\Users\RIMinullin\Documents\протоколы\фбу\word_files\scan5.docx')
    wp.get_data_by_keys()
    wp.get_indicators_from_word_file()
    # print(wp.data)
    wp.get_all_required_data_from_word_file()
    print(wp.data)
