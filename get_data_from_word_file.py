"""
Модуль для изъятия необходимых данных их word файла.

    class WordFileParser - содержит методы изъятия нужных данных из Word файла.
      Для получения необходимых данных при инициализации класса передать путь к word файлу.
      Для получения и возвращения словаря с необходимыми данными у экземпляра класса
      вызвать метод get_all_required_data_from_word_file без передачи аргументов.

      def get_all_required_data_from_word_file() - возвращает словарь, ключи необходимые для
        сбора поля, значения - сами данные из файла. Значения в словаре в виде строк.

      пример использования:
        word_parser = WordFileParser(path_to_word_file)
        required_data = word_parser.get_all_required_data_from_word_file()

"""
import re

from docx import Document
import docx2txt

from constants import REQUIRED_KEYS_FOR_PARSING_FIRST_PAGE
from comparison_results_and_norms import compare_result_and_norms

# Множество для хранения заголовков таблицы, для пропуска при высчитывании соответствия нормам.
KEYS_IN_INDICATORS_FOR_PASS = {'Токсичные элементы, мг/кг:',
            'Пестициды, мг/кг:', }


class WordFileParser:
    """ Класс для изъятия данных из Word файла."""
    def __init__(self, input_word_file: str):
        self.word_file = input_word_file
        self.all_text_from_file = self.get_one_string_from_word_file()
        self.document_with_tables = self.convert_word_file_to_docx_document()

    def get_one_string_from_word_file(self) -> str:
        """ Прочитать word файл, вернуть одну большую строку. """
        text = docx2txt.process(self.word_file)
        return text

    def convert_word_file_to_docx_document(self) -> Document:
        """ Прочитать word файл, преобразовать для работы с таблицами. """
        document = Document(self.word_file)
        return document

    def _find_substring_with_number_and_date(self):
        """Найти и вернуть подстроку с номером и датой протокола"""
        index_start_protocol = self.all_text_from_file.find('Протокол исп')
        index_end_protocol = index_start_protocol + self.all_text_from_file[index_start_protocol:].find('.')
        name_and_date_of_protocol = self.all_text_from_file[index_start_protocol:index_end_protocol]
        name_and_date_of_protocol = name_and_date_of_protocol.replace('\t', '').replace('\n', '')
        return name_and_date_of_protocol

    def _extract_protocol_date(self, text: str) -> str:
        """ Извлечь и вернуть из текста дату протокола. """
        # Сохраняем в список индексы цифр в тексте(дате протокола).
        indexes_of_digits_in_date = [index for index, symbol in enumerate(text) if symbol.isdigit()]

        if indexes_of_digits_in_date:
            # Вырезаем строку с датой и убираем все лишнее.
            date = text[indexes_of_digits_in_date[0]:indexes_of_digits_in_date[-1] + 1]
            cleaned_date = date.replace('»', '').replace(' ', '')
            formatted_date = add_spaces_between_digits_and_letters(cleaned_date)
            return formatted_date

    def extract_number_and_date_protocol(self) -> tuple[str, str]:
        """ Извлечь и вернуть из текста номер и дату протокола. """
        text_with_number_and_date = self._find_substring_with_number_and_date()

        number_pattern = r'№\s?(\d+)\s*\/\d+-Д'
        number_protocol = re.search(number_pattern, text_with_number_and_date)

        if number_protocol:
            text_with_date = text_with_number_and_date[number_protocol.end():]
            date_protocol = self._extract_protocol_date(text_with_date)
            number_protocol = number_protocol.group().replace(' ', '').replace('\t', '')
            return number_protocol, date_protocol
        else:
            return 'Номер протокола не найден', 'Дата протокола не найдена'

    def get_all_items_from_tables(self) -> dict:
        """ Собрать в словарь, все данные из таблиц word файла. """
        all_tables_in_file: list = self.document_with_tables.tables
        # Пустой словарь под данные из таблиц, количество ключей равно количеству таблиц в файле.
        data_by_tables: dict = {table: None for table in range(len(all_tables_in_file))}
        
        # Цикл по таблицам из Word файла.
        for number_of_table, table in enumerate(all_tables_in_file):
            # В словаре ключу - номеру таблицы создаем список списков для заполнения строками.
            data_by_tables[number_of_table] = [[] for _ in range(len(table.rows))]
            # Цикл по строкам одной таблицы из WORD файла.
            for number_of_row, row in enumerate(table.rows):
                # Цикл по ячейкам строки в таблице
                for cell in row.cells:
                    data_by_tables[number_of_table][number_of_row].append(cell.text)
        
        # Наполняем итоговый словарь данными из таблиц, исходим из структуры в виде двух колонок.
        data_from_tables = {}
        for item in data_by_tables:
            data_from_tables.update({i[0]: i[1] for i in data_by_tables[item]})

        return data_from_tables

    def leave_required_items_from_tables(
            self, required_keys: tuple = REQUIRED_KEYS_FOR_PARSING_FIRST_PAGE) -> dict:
        """ В словаре данных, из таблиц word файла оставить только нужные ключи.
        Без номера, дата протокола, и показателей. """

        def _find_code_of_store(data: dict) -> dict:
            """ Подфункция - найти код магазина - место отбора проб. """
            pattern_for_code = r'\d{5}'
            code = re.search(pattern_for_code, data['Место отбора проб'])
            data['Место отбора проб'] = code.group()
            return data

        data_from_tables: dict = self.get_all_items_from_tables()
        
        required_data_from_tables = {}
        # В выходной словарь добавляем только пары с нужными ключами.
        for item in data_from_tables.items():
            if item[0] in set(required_keys):
                required_data_from_tables[item[0]] = item[1]
        required_data_from_tables = _find_code_of_store(required_data_from_tables)
        return required_data_from_tables

    def check_that_table_consist_indicators(self, first_row_cells):
        if (len(first_row_cells) > 1 and
                (re.search(r"Наименование", first_row_cells[0].text)) and
                (re.search(r"Результат", first_row_cells[1].text)) and
                (re.search(r"Требования", first_row_cells[2].text))):
            return True

    def get_indicators_from_word_file(self) -> list:
        """ Получить показатели из протокола.

        Берет таблицы из документа, ищет в которых первые
        две ячейки в первой строке это "Наименование показателя" и "Результат".
        Если такие таблицы найдены, то значения из них сохраняются в словарь.
        Если таких таблиц несколько, то словари объединяются в список. """

        all_tables_in_file: list = self.document_with_tables.tables
        result_list = []

        for table in all_tables_in_file:  # Цикл по таблицам.

            if len(table.rows) > 1:  # Проверка на наличие в таблице более двух строк.
                first_row_cells = table.row_cells(0)

                # Проверка, что таблица, относится к показателям.
                if self.check_that_table_consist_indicators(first_row_cells):
                    indicators = {}
                    # Цикл по строкам, начиная со второй.
                    for _, row in enumerate(table.rows[1:], start=1):

                        cells = row.cells
                        indicator_name = cells[0].text.strip()  # Имя показателя
                        results = cells[1].text.strip()  # Результаты
                        requirements = cells[2].text.strip()  # Требования НД
                        indicators[indicator_name] = (results, requirements)

                    result_list.append(indicators)
        return result_list

    def add_a_conclusion_about_compliance_of_norms(self, data_from_protocol: dict):
        """ Добавить в словарь вывод о соответствие результатов исследований нормам."""
        # Переменная есть ли нарушения в показателях.
        violations = False
        # Перебираем словари в 'показателях'.
        for number, indicator_dict in enumerate(data_from_protocol['Показатели']):
            # key - наименование показателя, value - tuple, где value[0] - результат,
            # value[1] - норма для показателя
            for key, value in indicator_dict.items():
                if key in KEYS_IN_INDICATORS_FOR_PASS:
                    new_value = value + ('-',)
                else:
                    # Получаем bool соответствует ли результат норме.
                    complies_with_standards = compare_result_and_norms(value[0], value[1])
                    if not complies_with_standards:
                        violations = True
                    new_value = value + (complies_with_standards,)

                # Добавляем в словарь в value bool как вывод о соответствии норм.
                data_from_protocol['Показатели'][number][key] = new_value

        # Добавляем в словарь новый ключ с bool есть ли в протоколе нарушение норм.
        data_from_protocol['Нарушения норм'] = 'Не соответствуют' if violations else 'Соответствуют'
        return data_from_protocol

    def get_all_required_data_from_word_file(self) -> dict:
        """ Основной метод - направленный на изъятие данных из Word файла."""
        data_from_protocol = {}
        (data_from_protocol['Номер протокола'],
         data_from_protocol['Дата протокола']) = self.extract_number_and_date_protocol()

        # Добавить в словарь все пары с первой таблицы.
        data_from_protocol.update(self.leave_required_items_from_tables())

        # Добавить в словарь показатели.
        data_from_protocol['Показатели'] = self.get_indicators_from_word_file()
        data_from_protocol = self.add_a_conclusion_about_compliance_of_norms(data_from_protocol)
        # data_from_protocol['Показатели'] = json.dumps(data_from_protocol['Показатели'])
        return data_from_protocol


def add_spaces_between_digits_and_letters(text: str) -> str:
    """ Добавить пробелы между цифрами и месяцем. """
    result = ''
    for i in range(len(text) - 1):
        if text[i].isalpha() and text[i - 1].isdigit():
            result += (' ' + text[i])
        elif text[i].isalpha() and text[i + 1].isdigit():
            result += (text[i] + ' ')
        else:
            result += text[i]
    result += text[-1]
    return result
