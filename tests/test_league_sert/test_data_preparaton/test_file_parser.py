""" Тестирование модуля file_parser. """
import pytest

from docx import Document

from league_sert.data_preparation.exceptions import TypeOfTableError
from league_sert.data_preparation.file_parser import (
    get_text_with_superscripts, find_out_table_type,
    collect_prod_control_data, check_table_only_symbols)
from league_sert.data_preparation.extract_numb_and_date import get_main_numb_and_date
from tests.test_league_sert.test_data_preparaton.constants import TEST_WORD_FILES
from tests.test_league_sert.test_data_preparaton.data_for_test.test_data_text_numb_date import DATA_NUMB_DATE, \
    PROD_CONTROL_STR

document = Document()


def create_cell_with_text(doc, text: str, superscript_indices: list):
    """ Создать объект cell, с текстом, где будет содержаться степень. """
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    paragraph = cell.paragraphs[0]
    for index, char in enumerate(text):
        run = paragraph.add_run(char)
        if index in superscript_indices:
            run.font.superscript = True
    return cell


@pytest.mark.parametrize('cell, expected', [
    (create_cell_with_text(document, "5,5х102", superscript_indices=[6]), '5,5х10²'),
    (create_cell_with_text(document, "3,0x104", superscript_indices=[6]), '3,0x10⁴'),
    (create_cell_with_text(document, "102", superscript_indices=[2]), '10²'),
    (create_cell_with_text(document, "1030х102", superscript_indices=[2, 3, 7]), '10³⁰х10²'),
    (create_cell_with_text(document, "1030х102", superscript_indices=[2, 3, 7]), '10³⁰х10²'),
    (create_cell_with_text(document, "не обнаружены в 1 г (см3)",
                           superscript_indices=[23]), 'не обнаружены в 1 г (см³)'),

])
def test_get_text_with_superscripts(cell, expected):
    """ Сохранение степеней при чтении документа в регистре степени. """
    result = get_text_with_superscripts(cell)
    assert result == expected


def test_find_out_type_of_table_raises_error():
    """ Проверка возникновения исключения при тексте в первых
    ячейках таблицы, не подпадающем под паттерны."""
    cell0 = create_cell_with_text(document, "", superscript_indices=[])
    cell1 = create_cell_with_text(document, "Непонятные символы", superscript_indices=[])
    cell2 = create_cell_with_text(document, "some", superscript_indices=[])
    with pytest.raises(TypeOfTableError) as excinfo:
        find_out_table_type([cell0, cell1, cell2], document.tables[0])
    assert cell1.text in str(excinfo.value)


def test_get_main_numb_and_date():
    """ Протестировать извлечение номер и даты основной части протокола."""
    text = ("Акт отбора проб № 4639 от 13.05.2024 г. Количество зашифрованных проб "
            "\n1 Протокол составлен в 2-х экземплярах\n	Протокол испытаний "
            "№	4639 /24-Д от «29»			мая	20 24 г. Данные о пробе\n Шифр пробы")
    result = get_main_numb_and_date(text)
    assert isinstance(result, tuple)
    assert result[0] == '4639/24-Д'
    assert result[1] == '29 мая 2024'


@pytest.mark.parametrize('text, file, expect_numb, expect_date',
                         [(DATA_NUMB_DATE[0], TEST_WORD_FILES[0], '8281/24-Д', '29 августа 2024'),
                          (DATA_NUMB_DATE[1], TEST_WORD_FILES[1], '7241/24-Д', '25 июля 2024'),
                          (DATA_NUMB_DATE[2], TEST_WORD_FILES[2], '6789/24-Д', '22 июля 2024')])
def test_get_main_numb_and_date(text, file, expect_numb, expect_date):
    """ Протестировать извлечение номер и даты основной части протокола."""
    result = get_main_numb_and_date(text, file)
    assert isinstance(result, tuple)
    assert result[0] == expect_numb
    assert result[1] == expect_date


def test_collect_prod_control_data():
    """ Протестировать сбор данных по производственному контролю. """
    result_data = collect_prod_control_data(PROD_CONTROL_STR)
    assert result_data['number_of_protocol'] == '4218/24-Д'
    assert result_data['date_of_protocol'] == '07 мая 2024'
    assert result_data['act'] == '4218 ОТ 02.05.2024'
    assert result_data['place_of_measurement'] == (
        'Ленинградская обл., г. Гатчина, ул. Куприна, д.48А, лит. А, магазин № 47109')
    assert result_data['inner_conclusion'] == (
        '- на момент проведения измерений параметры микроклимата, '
        'уровни шума и общей вибрации соответствуют требованиям')


def test_check_table_only_symbols():
    """ Протестировать проверку таблиц. Задача - выявить
     таблицу, состоящую из одних символов."""

    doc = Document(TEST_WORD_FILES[3])
    table_ = doc.tables[0]
    only_symbols = check_table_only_symbols(table_)
    assert only_symbols is False

    table_ = doc.tables[1]
    only_symbols = check_table_only_symbols(table_)
    assert only_symbols is True

    table_ = doc.tables[2]
    only_symbols = check_table_only_symbols(table_)
    assert only_symbols is True
