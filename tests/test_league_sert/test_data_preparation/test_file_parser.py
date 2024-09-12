""" Тестирование модуля file_parser. """

import pytest

from docx import Document

from league_sert.data_preparation.exceptions import TypeOfTableError
from league_sert.data_preparation.file_parser import (
    get_text_with_superscripts, find_out_type_of_table,
    get_main_numb_and_date, collect_prod_control_data)

document = Document()


def create_cell_with_text(doc, text: str, superscript_indices:list):
    """ Создать объект cell, с текстом, где будет содержаться степень. """
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    paragraph = cell.paragraphs[0]
    for index, char in enumerate(text):
        run = paragraph.add_run(char)
        if index in superscript_indices:
            run.font.superscript = True
    return cell


@pytest.mark.parametrize('cell, expected', [
    (create_cell_with_text(document, "5,5х102", superscript_indices=[6]), '5,5х10²'),
    (create_cell_with_text(document, "3,0x104", superscript_indices=[6]), '3,0x10⁴'),
    (create_cell_with_text(document, "102", superscript_indices=[2]), '10²'),
    (create_cell_with_text(document, "1030х102", superscript_indices=[2, 3, 7]), '10³⁰х10²'),
    (create_cell_with_text(document, "1030х102", superscript_indices=[2, 3, 7]), '10³⁰х10²'),
    (create_cell_with_text(document, "не обнаружены в 1 г (см3)",
                           superscript_indices=[23]), 'не обнаружены в 1 г (см³)'),

])
def test_get_text_with_superscripts(cell, expected):
    """ Сохранение степеней при чтении документа в регистре степени. """
    result = get_text_with_superscripts(cell)
    assert result == expected


# def test_find_out_type_of_table_raises_error():
#     """ Проверка возникновения исключения при содержимом таблицы,
#     не подпадающем под паттерны."""
#     cell1 = create_cell_with_text(document, "", superscript_indices=[])
#     cell2 = create_cell_with_text(document, "", superscript_indices=[])
#     with pytest.raises(TypeOfTableError) as excinfo:
#         print('1!'*30, find_out_type_of_table([cell1, cell2]))
#
#     assert cell1 in str(excinfo.value)


def test_get_main_numb_and_date():
    """ Протестировать извлечение номер и даты основной части протокола."""
    text = ("Акт отбора проб № 4639 от 13.05.2024 г. Количество зашифрованных проб "
            "\n1 Протокол составлен в 2-х экземплярах\n	Протокол испытаний "
            "№	4639 /24-Д от «29»			мая	20 24 г. Данные о пробе\n Шифр пробы")
    result = get_main_numb_and_date(text)
    assert isinstance(result, tuple)
    assert result[0] == '4639/24-Д'
    assert result[1] == '29 мая 2024'


PROD_CONTROL_STR = (
    '\nАттестат аккредитации: РОСС RU.31112.21ПР52\nЛицензия: 77.01.13.001 .Л.000039.09.19'
    '\n\n\n\nПРОТОКОЛ № 4218/24-Д\nизмерений по производственному контролю\nот «07» мая 2024 '
    'г.\n\n\tНомер и дата заявки: 8/24 от 10.01.2024 г.\n\n\tЗаявитель и его адрес '
    '(юридический и фактический): АО «ДИКСИ-Юг», Московская обл., г. Подольск, ул. Юбилейная, '
    'д. 32А\n\n\t№ Акта и дата проведения измерений: 4218 ОТ 02.05.2024 г.\n\n\t'
    'Адрес проведения измерений: Ленинградская обл., г. Гатчина, ул. Куприна, д.48А, лит. А, '
    'магазин № 47109\n\n\tСведения о средствах измерения:\n\n\nК протоколу № 4218/24-Д'
    ' от 07.05.2024 г.\n\nЗАКЛЮЧЕНИЕ\n\nРезультаты проведенных измерений показали, '
    'что на объекте заказчика АО «ДИКСИ-ЮГ», расположенном по адресу: '
    'Ленинградская обл., г. Гатчина, ул. Куприна, Д.48А, лит. А, магазин № 47109\n\n- '
    'на момент проведения измерений параметры микроклимата, уровни шума и общей вибрации '
    'соответствуют требованиям\n\nСанПиН 1.2.3685-21 «Гигиенические нормативы и требования'
    ' к обеспечению безопасности и (или) безвредности для человека факторов среды обитания».'
    '\n\n\n\nСанитарный врач\n\nСанитарный врач\n\n(подпись)\n')


def test_collect_prod_control_data():
    """ Протестировать сбор данных по производственному контролю. """
    result_data = collect_prod_control_data(PROD_CONTROL_STR)
    assert result_data['number_of_protocol'] == '4218/24-Д'
    assert result_data['date_of_protocol'] == '07 мая 2024'
    assert result_data['act'] == '4218 ОТ 02.05.2024'
    assert result_data['place_of_measurement'] == (
        'Ленинградская обл., г. Гатчина, ул. Куприна, д.48А, лит. А, магазин № 47109')
    assert result_data['inner_conclusion'] == (
        '- на момент проведения измерений параметры микроклимата, '
        'уровни шума и общей вибрации соответствуют требованиям')
